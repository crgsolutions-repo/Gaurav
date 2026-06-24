from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SOURCE_DIR = ROOT / "policies" / "source"
OUTPUT_DIR = ROOT / "policies" / "generated"
OUTPUT_PATH = OUTPUT_DIR / "Enterprise_HR_Policy_Handbook_Demo.docx"
PDF_PATH = OUTPUT_DIR / "Enterprise_HR_Policy_Handbook_Demo.pdf"
BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
MUTED = RGBColor(89, 96, 105)


def set_cell_shading(paragraph, fill):
    properties = paragraph._p.get_or_add_pPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    properties.append(shading)


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Page ")
    run.font.size = Pt(9)
    run.font.color.rgb = MUTED
    field = OxmlElement("w:fldSimple")
    field.set(qn("w:instr"), "PAGE")
    paragraph._p.append(field)


def configure_document(document):
    section = document.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    section.different_first_page_header_footer = True

    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for style_name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, DARK_BLUE, 10, 5),
    ):
        style = styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    header = section.header.paragraphs[0]
    header.text = "Enterprise HR Policy Handbook | Demonstration"
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for run in header.runs:
        run.font.name = "Calibri"
        run.font.size = Pt(9)
        run.font.color.rgb = MUTED
    add_page_number(section.footer.paragraphs[0])


def add_cover(document):
    document.add_paragraph().paragraph_format.space_after = Pt(42)
    kicker = document.add_paragraph("HR ASSISTANT KNOWLEDGE BASE")
    kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER
    kicker.paragraph_format.space_after = Pt(14)
    run = kicker.runs[0]
    run.bold = True
    run.font.size = Pt(10)
    run.font.color.rgb = BLUE

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(10)
    run = title.add_run("Enterprise HR Policy Handbook")
    run.bold = True
    run.font.name = "Calibri"
    run.font.size = Pt(25)
    run.font.color.rgb = RGBColor(20, 32, 48)

    subtitle = document.add_paragraph("Demonstration policy corpus for conversational RAG testing")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(36)
    subtitle.runs[0].font.size = Pt(13)
    subtitle.runs[0].font.color.rgb = MUTED

    notice = document.add_paragraph()
    notice.paragraph_format.left_indent = Inches(0.45)
    notice.paragraph_format.right_indent = Inches(0.45)
    notice.paragraph_format.space_before = Pt(10)
    notice.paragraph_format.space_after = Pt(18)
    set_cell_shading(notice, "F2F4F7")
    run = notice.add_run(
        "Important: This handbook is demonstration content. It is not an adopted company policy and does not replace HR, legal, payroll, tax, safety, or regulatory review."
    )
    run.bold = True
    run.font.size = Pt(10.5)
    run.font.color.rgb = RGBColor(92, 48, 20)

    meta = document.add_paragraph("Version 1.0-demo | Prepared June 2026")
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.runs[0].font.size = Pt(10)
    meta.runs[0].font.color.rgb = MUTED
    document.add_page_break()


def markdown_parts(path):
    title = path.stem.replace("_", " ").title()
    parts = []
    paragraph_lines = []

    def flush():
        content = " ".join(line.strip() for line in paragraph_lines if line.strip()).strip()
        if content:
            parts.append(("paragraph", content))
        paragraph_lines.clear()

    for line in path.read_text(encoding="utf-8").splitlines():
        if line.startswith("# "):
            flush()
            title = line[2:].strip()
        elif line.startswith("## "):
            flush()
            parts.append(("heading", line[3:].strip()))
        elif line.lower().startswith(("category:", "version:")):
            flush()
            parts.append(("metadata", line.strip()))
        elif not line.strip():
            flush()
        else:
            paragraph_lines.append(line)
    flush()
    return title, parts


def add_policy(document, path, first=False):
    title, parts = markdown_parts(path)
    if not first:
        document.add_page_break()
    document.add_heading(title, level=1)
    for kind, value in parts:
        if kind == "heading":
            document.add_heading(value, level=2)
        elif kind == "metadata":
            paragraph = document.add_paragraph(value)
            paragraph.paragraph_format.space_after = Pt(2)
            paragraph.runs[0].font.size = Pt(9)
            paragraph.runs[0].font.color.rgb = MUTED
        else:
            paragraph = document.add_paragraph(value)
            paragraph.paragraph_format.widow_control = True


def main():
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    document = Document()
    configure_document(document)
    add_cover(document)

    intro = document.add_heading("How to Use This Demonstration Handbook", level=1)
    intro.paragraph_format.space_before = Pt(0)
    document.add_paragraph(
        "The HR Assistant retrieves relevant sections from this corpus and supplies them to Gemini together with authorized employee context. Responses should cite the policy title, distinguish advice from completed action, and state when the available policy is insufficient."
    )
    documents = sorted(SOURCE_DIR.glob("*.md"))
    document.add_heading("Included Policy Areas", level=2)
    for path in documents:
        title, _parts = markdown_parts(path)
        paragraph = document.add_paragraph(style="List Bullet")
        paragraph.add_run(title)

    for index, path in enumerate(documents):
        add_policy(document, path, first=False)

    document.add_page_break()
    document.add_heading("Governance Note", level=1)
    document.add_paragraph(
        "Before production use, each policy must have an owner, effective date, review date, jurisdiction, approved reporting contacts, configured authority limits, and formal publication status. Superseded versions should be archived rather than overwritten so retrieval and audit history remain explainable."
    )
    document.core_properties.title = "Enterprise HR Policy Handbook - Demonstration"
    document.core_properties.subject = "RAG policy corpus for the HR Assistant project"
    document.core_properties.author = "HR Assistant Project"
    document.save(OUTPUT_PATH)
    print(OUTPUT_PATH)
    build_pdf(documents)
    print(PDF_PATH)


def build_pdf(documents):
    from reportlab.lib import colors
    from reportlab.lib.enums import TA_CENTER
    from reportlab.lib.pagesizes import letter
    from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
    from reportlab.lib.units import inch
    from reportlab.platypus import CondPageBreak, PageBreak, Paragraph, SimpleDocTemplate, Spacer

    styles = getSampleStyleSheet()
    body = ParagraphStyle(
        "PolicyBody",
        parent=styles["BodyText"],
        fontName="Helvetica",
        fontSize=10.5,
        leading=13.2,
        spaceAfter=6,
        textColor=colors.HexColor("#1F2933"),
    )
    title_style = ParagraphStyle(
        "CoverTitle",
        parent=styles["Title"],
        fontName="Helvetica-Bold",
        fontSize=25,
        leading=30,
        alignment=TA_CENTER,
        textColor=colors.HexColor("#142030"),
        spaceAfter=10,
    )
    subtitle_style = ParagraphStyle(
        "CoverSubtitle",
        parent=body,
        fontSize=13,
        leading=17,
        alignment=TA_CENTER,
        textColor=colors.HexColor("#596069"),
        spaceAfter=30,
    )
    h1 = ParagraphStyle(
        "PolicyH1",
        parent=styles["Heading1"],
        fontName="Helvetica-Bold",
        fontSize=16,
        leading=20,
        textColor=colors.HexColor("#2E74B5"),
        spaceBefore=18,
        spaceAfter=10,
    )
    h2 = ParagraphStyle(
        "PolicyH2",
        parent=styles["Heading2"],
        fontName="Helvetica-Bold",
        fontSize=13,
        leading=16,
        textColor=colors.HexColor("#2E74B5"),
        spaceBefore=14,
        spaceAfter=7,
        keepWithNext=True,
    )
    metadata = ParagraphStyle(
        "Metadata",
        parent=body,
        fontSize=9,
        leading=11,
        textColor=colors.HexColor("#596069"),
        spaceAfter=2,
    )
    notice = ParagraphStyle(
        "Notice",
        parent=body,
        fontName="Helvetica-Bold",
        fontSize=10.5,
        leading=14,
        textColor=colors.HexColor("#5C3014"),
        backColor=colors.HexColor("#F2F4F7"),
        borderPadding=12,
        spaceBefore=10,
        spaceAfter=18,
    )

    def footer(canvas, doc):
        canvas.saveState()
        canvas.setFont("Helvetica", 8.5)
        canvas.setFillColor(colors.HexColor("#596069"))
        if doc.page > 1:
            canvas.drawString(inch, 0.55 * inch, "Enterprise HR Policy Handbook | Demonstration")
        canvas.drawRightString(7.5 * inch, 0.55 * inch, f"Page {doc.page}")
        canvas.restoreState()

    story = [
        Spacer(1, 0.75 * inch),
        Paragraph("HR ASSISTANT KNOWLEDGE BASE", metadata),
        Spacer(1, 0.12 * inch),
        Paragraph("Enterprise HR Policy Handbook", title_style),
        Paragraph("Demonstration policy corpus for conversational RAG testing", subtitle_style),
        Paragraph(
            "Important: This handbook is demonstration content. It is not an adopted company policy and does not replace HR, legal, payroll, tax, safety, or regulatory review.",
            notice,
        ),
        Paragraph("Version 1.0-demo | Prepared June 2026", subtitle_style),
        PageBreak(),
        Paragraph("How to Use This Demonstration Handbook", h1),
        Paragraph(
            "The HR Assistant retrieves relevant sections from this corpus and supplies them to Gemini together with authorized employee context. Responses should cite the policy title, distinguish advice from completed action, and state when the available policy is insufficient.",
            body,
        ),
        Paragraph("Included Policy Areas", h2),
    ]
    for path in documents:
        policy_title, _parts = markdown_parts(path)
        story.append(Paragraph(policy_title, body, bulletText="-"))

    for path in documents:
        policy_title, parts = markdown_parts(path)
        story.extend([CondPageBreak(4 * inch), Paragraph(policy_title, h1)])
        for kind, value in parts:
            if kind == "heading":
                story.append(Paragraph(value, h2))
            elif kind == "metadata":
                story.append(Paragraph(value, metadata))
            else:
                story.append(Paragraph(value, body))

    story.extend(
        [
            PageBreak(),
            Paragraph("Governance Note", h1),
            Paragraph(
                "Before production use, each policy must have an owner, effective date, review date, jurisdiction, approved reporting contacts, configured authority limits, and formal publication status. Superseded versions should be archived rather than overwritten so retrieval and audit history remain explainable.",
                body,
            ),
        ]
    )
    pdf = SimpleDocTemplate(
        str(PDF_PATH),
        pagesize=letter,
        rightMargin=inch,
        leftMargin=inch,
        topMargin=inch,
        bottomMargin=0.8 * inch,
        title="Enterprise HR Policy Handbook - Demonstration",
        author="HR Assistant Project",
    )
    pdf.build(story, onFirstPage=footer, onLaterPages=footer)


if __name__ == "__main__":
    main()
